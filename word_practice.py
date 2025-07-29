# -*- coding: utf-8 -*-
import sqlite3
import random
import argparse
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
<<<<<<< HEAD
import os
=======
>>>>>>> 8e7781ddd5932940ce6300496ab4a8827ce32409

def fetch_words(db_path, table, count):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
<<<<<<< HEAD
    # 检查表结构
    cursor.execute(f"PRAGMA table_info({table})")
    columns = [row[1] for row in cursor.fetchall()]
    select_cols = []
    if 'en' in columns:
        select_cols.append('en')
    if 'pro' in columns:
        select_cols.append('pro')
    if 'zh' in columns:
        select_cols.append('zh')
    if 'promt' in columns:
        select_cols.append('promt')
    select_clause = ', '.join(select_cols)
    cursor.execute(f'SELECT {select_clause} FROM {table}')
    rows = cursor.fetchall()
    conn.close()
    # 生成序号（1-based）
    rows_with_index = [(i+1,) + row for i, row in enumerate(rows)]
    if count == -1 or len(rows_with_index) <= count:
        print(f"提取表 {table} 的所有 {len(rows_with_index)} 条记录")
        return rows_with_index
    else:
        sampled = random.sample(rows_with_index, count)
        sampled.sort(key=lambda x: x[0])  # 保证序号递增
        return sampled
=======
    cursor.execute(f"SELECT english, chinese FROM {table}")
    rows = cursor.fetchall()
    conn.close()
    
    if count == -1:
        # 提取所有单词
        print(f"提取表 {table} 的所有 {len(rows)} 条记录")
        return rows
    elif len(rows) < count:
        print(f"警告：表 {table} 仅有 {len(rows)} 条，已全部抽取。")
        return rows
    else:
        return random.sample(rows, count)
>>>>>>> 8e7781ddd5932940ce6300496ab4a8827ce32409

def make_underline(text, ratio=1.5, keep_semicolon=False):
    result = []
    for ch in text:
        if keep_semicolon and ch in ['；', ';']:
            result.append(ch)
        elif ch.strip() == '':
            result.append(' ')
        else:
            n = max(1, int(round(len(ch) * ratio)))
            result.append('_' * n)
    return ''.join(result)



<<<<<<< HEAD
def generate_docx(words, mode, output_path, shuffled_words=None, answer_mode=False):
    doc = Document()
    doc.add_heading('单词填词表' if not answer_mode else '单词答案表', 0)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '英文'
    hdr_cells[2].text = '词性'
    if answer_mode:
        hdr_cells[3].text = '中文'
    elif mode == 'chinese':
        hdr_cells[3].text = '中文+提示'
    elif mode == 'english':
        hdr_cells[3].text = '中文'
    else:
        hdr_cells[3].text = '中文+提示'
    for cell in hdr_cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
=======
def generate_docx(words, mode, output_path, shuffled_words=None):
    doc = Document()
    doc.add_heading('单词填词表', 0)
    table = doc.add_table(rows=1, cols=3)
    
    # 设置表格样式
    table.style = 'Table Grid'
    
    # 设置表头并垂直居中
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '英文'
    hdr_cells[1].text = '词性'
    hdr_cells[2].text = '中文+提示'
    for cell in hdr_cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # 使用传入的随机序列或生成新的
>>>>>>> 8e7781ddd5932940ce6300496ab4a8827ce32409
    if shuffled_words is None:
        shuffled_words = list(words)
        random.shuffle(shuffled_words)
    
    for item in shuffled_words:
<<<<<<< HEAD
        idx = item[0] if len(item) > 0 else ''
        eng = item[1] if len(item) > 1 else ''
        pro = item[2] if len(item) > 2 else ''
        chn = item[3] if len(item) > 3 else ''
        promt = item[4] if len(item) > 4 else ''
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx) if idx else ''
        if answer_mode:
            row_cells[1].text = str(eng) if eng else ''
        elif mode == 'chinese':
            space_len = int(len(str(eng)) * 1.5)
            row_cells[1].text = ' ' * space_len if eng else ''
        else:
            row_cells[1].text = str(eng) if eng else ''
        row_cells[2].text = str(pro) if pro else ''
        if answer_mode:
            row_cells[3].text = str(chn) if chn else ''
        elif mode == 'chinese' or mode == 'both':
            chn_text = str(chn) if chn else ''
            if promt:
                chn_text += f'（{promt}）'
            row_cells[3].text = chn_text
        elif mode == 'english':
            row_cells[3].text = str(chn) if chn else ''
        else:
            chn_text = str(chn) if chn else ''
            if promt:
                chn_text += f'（{promt}）'
            row_cells[3].text = chn_text
        for cell in row_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    doc.save(output_path)
    if answer_mode:
        print(f"已生成答案表：{output_path}")
    else:
        print(f"已生成填词表：{output_path}")
=======
        # 兼容不同数据结构
        if isinstance(item, dict):
            eng = item.get('english') or item.get('en') or ''
            pro = item.get('pos') or item.get('pro') or ''
            chn = item.get('chinese') or item.get('zh') or ''
            promt = item.get('promt') or ''
        else:
            # 兼容原有tuple结构
            if len(item) == 2:
                eng, chn = item
                pro = ''
                promt = ''
            elif len(item) == 3:
                eng, pro, chn = item
                promt = ''
            elif len(item) == 4:
                eng, pro, chn, promt = item
            else:
                eng = chn = pro = promt = ''
        if eng is None and chn is None:
            continue
        row_cells = table.add_row().cells
        for cell in row_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row_cells[0].text = str(eng) if eng else ''
        row_cells[1].text = str(pro) if pro else ''
        # 中文+promt
        chn_text = str(chn) if chn else ''
        if promt:
            chn_text += f'（{promt}）'
        row_cells[2].text = chn_text
    doc.save(output_path)
    print(f"已生成填词表（随机排序）：{output_path}")
>>>>>>> 8e7781ddd5932940ce6300496ab4a8827ce32409
    return shuffled_words

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='单词/词组抽取与填词表生成')
<<<<<<< HEAD
    parser.add_argument('--db', type=str, help='sqlite db文件')
    parser.add_argument('--letter', type=str, nargs='+', required=True, help='首字母表（如a、b、c，可多个）')
    parser.add_argument('--count', type=int, default=10, help='抽取数量（输入-1提取所有单词）')
    parser.add_argument('--mode', type=str, choices=['chinese', 'english', 'both'], nargs='+', default=['chinese'], help='填词表模式（可多个）')
    parser.add_argument('--output', type=str, nargs='+', help='输出docx文件名（可多个，与mode一一对应，不指定则自动命名）')
    args = parser.parse_args()

    db_path = args.db
    while not db_path:
        db_path = input('请输入sqlite db文件路径: ').strip()

    for letter in args.letter:
        table = letter.lower()
        words = fetch_words(db_path, table, args.count)
    if not words:
        print(f"未找到任何数据，请检查表名和数据库内容。({table})")
        # 自动生成输出文件名
    outputs = args.output if args.output and len(args.output) == len(args.mode) else [f"{table}_{mode}.docx" for mode in args.mode]
    if len(outputs) != len(args.mode):
        print(f"错误：mode参数数量({len(args.mode)})与output参数数量({len(outputs)})不匹配")
        # 每个首字母独立随机序列
    shuffled_words = list(words)
    random.shuffle(shuffled_words)
        # 确保PracticeFile文件夹存在
    out_dir = "PracticeFile"
    if not os.path.exists(out_dir):
        os.makedirs(out_dir)
    for i, (mode, output) in enumerate(zip(args.mode, outputs)):
        out_path = os.path.join(out_dir, output)
        print(f"\n生成{table}第{i+1}个文件：模式={mode}, 输出={out_path}")
        generate_docx(words, mode, out_path, shuffled_words)
        # 生成key答案文件
    key_output = os.path.join(out_dir, f"{table}_key.docx")
    print(f"\n生成答案文件：{key_output}")
    generate_docx(words, 'both', key_output, shuffled_words, answer_mode=True) 
=======
    parser.add_argument('--db', type=str, default='words.db', help='sqlite db文件')
    parser.add_argument('--letter', type=str, required=True, help='首字母表（如a、b、c）')
    parser.add_argument('--count', type=int, default=10, help='抽取数量（输入-1提取所有单词）')
    parser.add_argument('--mode', type=str, choices=['chinese', 'english', 'both'], nargs='+', default=['chinese'], help='填词表模式（可多个）')
    parser.add_argument('--output', type=str, nargs='+', default=['fill_table.docx'], help='输出docx文件名（可多个，与mode一一对应）')
    args = parser.parse_args()

    table = args.letter.lower()
    words = fetch_words(args.db, table, args.count)
    if not words:
        print(f"未找到任何数据，请检查表名和数据库内容。")
        exit(1)
    
    # 确保mode和output数量一致
    if len(args.mode) != len(args.output):
        print(f"错误：mode参数数量({len(args.mode)})与output参数数量({len(args.output)})不匹配")
        exit(1)
    
    # 生成相同的随机序列
    shuffled_words = list(words)
    random.shuffle(shuffled_words)
    
    # 为每个mode-output对生成docx文件
    for i, (mode, output) in enumerate(zip(args.mode, args.output)):
        print(f"\n生成第{i+1}个文件：模式={mode}, 输出={output}")
        generate_docx(words, mode, output, shuffled_words) 
>>>>>>> 8e7781ddd5932940ce6300496ab4a8827ce32409
